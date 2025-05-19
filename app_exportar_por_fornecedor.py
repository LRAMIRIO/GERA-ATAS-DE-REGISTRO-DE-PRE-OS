
import streamlit as st
import pandas as pd
import os
from io import BytesIO
from zipfile import ZipFile
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Side
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

st.title("Gerar planilhas e documentos Word por fornecedor")
st.markdown("Envie uma planilha preenchida (.xlsx) para gerar arquivos Excel e Word para cada fornecedor.")

uploaded_file = st.file_uploader("📂 Envie a planilha preenchida (.xlsx):", type="xlsx")

if uploaded_file:
    df = pd.read_excel(uploaded_file, skiprows=2)
    df.columns = [col.strip() for col in df.columns]

    borda_fina = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    with tempfile.TemporaryDirectory() as temp_dir:
        for fornecedor, grupo in df.groupby("FORNECEDOR"):
            nome_limpo = fornecedor[:40].replace(" ", "_").replace("/", "-")
            grupo["ITEM"] = grupo["ITEM"].astype("Int64")
            grupo_formatado = grupo[[
                "ITEM", "DESCRIÇÃO DO MATERIAL", "MARCA", "UNIDADE",
                "QUANTIDADE", "VALOR UNITÁRIO", "VALOR TOTAL"
            ]]

            # Excel
            wb = Workbook()
            ws = wb.active
            ws.title = "Itens Vencedores"
            ws.append(grupo_formatado.columns.tolist())

            for _, row in grupo_formatado.iterrows():
                ws.append(row.tolist())

            for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
                for cell in row:
                    if cell.row == 1:
                        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    elif cell.column_letter == "B":
                        cell.alignment = Alignment(horizontal="justify", vertical="center", wrap_text=True)
                    else:
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border = borda_fina

            for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
                row[5].number_format = 'R$ #,##0.00'
                row[6].number_format = 'R$ #,##0.00'

            ws.column_dimensions["A"].width = 6
            ws.column_dimensions["B"].width = 65
            ws.column_dimensions["C"].width = 20
            ws.column_dimensions["D"].width = 20
            ws.column_dimensions["E"].width = 12
            ws.column_dimensions["F"].width = 16
            ws.column_dimensions["G"].width = 16

            excel_path = os.path.join(temp_dir, f"{nome_limpo}.xlsx")
            wb.save(excel_path)

            # Word
            doc = Document()
            doc.add_heading(f'Fornecedor: {fornecedor}', 0)

            table = doc.add_table(rows=1, cols=len(grupo_formatado.columns))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            for i, col in enumerate(grupo_formatado.columns):
                hdr_cells[i].text = str(col)
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for _, row in grupo_formatado.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    texto = str(value).replace(" ", "\n") if grupo_formatado.columns[i] == "UNIDADE" else str(value)
                    row_cells[i].text = texto
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            docx_path = os.path.join(temp_dir, f"{nome_limpo}.docx")
            doc.save(docx_path)

        # Compactar os arquivos
        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, "w") as zipf:
            for f in os.listdir(temp_dir):
                caminho_completo = os.path.join(temp_dir, f)
                zipf.write(caminho_completo, arcname=f)

        st.success("✅ Arquivos gerados com sucesso!")
        st.download_button("📦 Baixar arquivos em .zip", zip_buffer.getvalue(), file_name="planilhas_e_docx_por_fornecedor.zip", mime="application/zip")
